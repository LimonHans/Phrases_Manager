from tkinter import *
from docx import Document
from docx.oxml.ns import qn
from docx.shared import Pt
import linecache as f_cache
import random as rand
import os

counter = 0

#################################################
#This 'Source_Code.py' contains no explanations.#
#If necessary, please open 'Explained_Code.py'. #
#################################################

def text_adder(fir, sec, mark):
    if fir == "" or sec == "":
        print("No Words Detected")
        return
    sentence = fir + ", " + sec
    print("\"%s\" will be written" % sentence)
    sentence += "\n"
    if mark != "":
        print("marked")
        with open("Corder/" + mark.lower() + ".marked_word", "a", encoding = 'utf-8') as goal: goal.write(sentence)
    else:
        print("No marks")
        with open("Corder/" + fir[0].lower() + ".common_word", "a", encoding = 'utf-8') as goal: goal.write(sentence)
    return
def Phrases_Adder():
    Adder_Window = Tk(baseName = "1", className = "1")
    Adder_Window.title("Phrases Adder")
    Adder_Window.resizable(False, False)
    Adder_Window.geometry("+620+410")

    Label(Adder_Window, text = "短语").grid(row = 0)
    Label(Adder_Window, text = "释义").grid(row = 1)
    Label(Adder_Window, text = "特殊标记").grid(row = 2)
    entry1 = Entry(Adder_Window)
    entry2 = Entry(Adder_Window)
    entry3 = Entry(Adder_Window)
    entry1.grid(row = 0, column = 1, padx = 15, pady = 5)
    entry2.grid(row = 1, column = 1, padx = 15, pady = 5)
    entry3.grid(row = 2, column = 1, padx = 15, pady = 5)

    def passer(func):
        text_us = entry1.get()
        text_zh = entry2.get()
        mark = entry3.get()
        text_adder(text_us, text_zh, mark)
        entry1.delete(0, "end")
        entry2.delete(0, "end")
        entry3.delete(0, "end")
        if func: Adder_Window.destroy()
    Button(Adder_Window, text = "获取并继续", width = 20, command = lambda: passer(False)).grid(row = 3, column = 0, sticky = "w", padx = 10, pady = 5)
    Button(Adder_Window, text = "获取并退出", width = 20, command = lambda: passer(True)).grid(row = 3, column = 1, sticky = "e", padx = 10, pady = 5)

    Adder_Window.mainloop()
    return
def Files_finder():
    ans = []
    for root, dirs, files in os.walk("Corder/"):
        for file in files:
            ans.append("Corder/" + file)
    return ans
def Text_finder(place):
    ans = f_cache.getlines(place)
    return ans[rand.randint(0, ans.__len__() - 1)].split(", ", 1)
def Phrases_Tester():
    list = Files_finder()
    Tester_Window = Tk(baseName = "2", className = "2")
    Tester_Window.title("Phrases Tester")
    Tester_Window.resizable(False, False)
    Tester_Window.geometry("+600+400")

    alpha = [0, StringVar(), StringVar(), "", StringVar()]
    alpha[4].set("开始测试")

    def FollowText(alpha):
        #print("Receive: " + alpha[0].__str__() + alpha[1].__str__() + alpha[2].__str__())
        if alpha[0] == 0:
            Ustext, Zhtext = Text_finder(list[rand.randint(0, list.__len__() - 1)])
            Zhtext = Zhtext[:Zhtext.__len__() - 1]
            print("Now testing " + Ustext + ", " + Zhtext)
            alpha[1].set(Ustext)
            alpha[2].set("")
            alpha[3] = Zhtext
            alpha[4].set("展示释义")
        elif alpha[0] == 1:
            alpha[2].set(alpha[3])
            alpha[4].set("下一组短语")
        alpha[0] = (alpha[0] + 1)%2
        #print("Out: " + alpha[0].__str__() + alpha[1].__str__() + alpha[2].__str__())
    Uslabel = Label(Tester_Window, textvariable = alpha[1], width = 20, height = 4)
    Uslabel.grid(row = 0, column = 0, columnspan = 2)
    Zhlabel = Label(Tester_Window, textvariable = alpha[2], width = 20, height = 4)
    Zhlabel.grid(row = 1, column = 0)
    Changer = Button(Tester_Window, textvariable = alpha[4], width = 20, height = 4, command = lambda: FollowText(alpha))
    Changer.grid(row = 1, column = 1, sticky = "e", padx = 10, pady = 5)

    Tester_Window.mainloop()
    return
def Sort_phrases(place):
    goal = f_cache.getlines(place)
    goal.sort()
    f = open(place, "w", encoding = "utf-8")
    for sentence in goal: f.write(sentence)
    return
def Shuffle_phrases(place):
    goal = f_cache.getlines(place)
    rand.shuffle(goal)
    f = open(place, "w", encoding = "utf-8")
    for sentence in goal: f.write(sentence)
    return
def Asker(Title, Question):
    ans = [False]
    Asker_Window = Tk(baseName = "3", className = "3")
    Asker_Window.title(Title)
    Asker_Window.resizable(False, False)
    Asker_Window.geometry("+600+420")

    def Answer(Value): ans[0] = Value
    Label(Asker_Window, text = Question, width = 40, height = 4).grid(row = 0, column = 0, columnspan = 2)
    Button(Asker_Window, text = "是", width = 20, height = 2, command = lambda: {Asker_Window.destroy(), Answer(True)}).grid(row = 1, column = 0)
    Button(Asker_Window, text = "否", width = 20, height = 2, command = lambda: {Asker_Window.destroy(), Answer(False)}).grid(row = 1, column = 1)

    Asker_Window.mainloop()
    return ans[0]
def Noticer(Title, Result):
    Noticer_Window = Tk(baseName = "4", className = "4")
    Noticer_Window.title(Title)
    Noticer_Window.resizable(False, False)
    Noticer_Window.geometry("+600+420")

    Label(Noticer_Window, text = Result, width = 40, height = 4).grid(row = 0, column = 0)
    Button(Noticer_Window, text = "确认", width = 40, height = 2, command = Noticer_Window.destroy).grid(row = 1, column = 0)

    Noticer_Window.mainloop()
    return
def Exporter(mode):
    if not Asker("Phrases Exporter", "是否确认导出？"): return
    list = Files_finder()
    for files in list:
        if mode: Sort_phrases(files)
        else: Shuffle_phrases(files)

    if mode: func_word = "顺序"
    else: func_word = "乱序"
    doc = Document()
    counter = "1"
    with open("0config", "r", encoding = "utf-8") as f:
        counter = f.read()
        print(counter + " read!")
        doc.add_heading("短语总结[第%s版 %s]" %(counter, func_word), 0)
    with open("0config", "w", encoding = "utf-8") as f: f.write(str(int(counter) + 1))
    #par.add_run('This is bold.').bold = True
    #par.add_run('\n This is italic.').italic = True
    if mode:
        for i in range(0, list.__len__()):
            if list[i].find(".common_word") != -1: list[i] = list[i].replace("Corder/", "0").replace(".common_word", "")
            else: list[i] = list[i].replace("Corder/", "1").replace(".marked_word", "")
        list.sort()
        #for i in range(0, list.__len__()): print(list[i])
        #return
    else: rand.shuffle(list)
    for files in list:
        if mode:
            if files[0] == "0": place = "Corder/" + files[1:] + ".common_word"
            else: place = "Corder/" + files[1:] + ".marked_word"
            now_part = f_cache.getlines(place)
            file_name = files[1:]
        else:
            now_part = f_cache.getlines(files)
            file_name = files.replace("Corder/", "").replace(".common_word", "").replace(".marked_word", "")
        for i in range(0, now_part.__len__()): now_part[i] = now_part[i][:now_part[i].__len__() - 1]
        doc.add_heading("From %s:" %file_name, 3)
        for sentence in now_part:
            Usword, Zhword = sentence.split(", ", 1)
            now_paragraph = doc.add_paragraph()
            Usrun = now_paragraph.add_run(Usword)
            Usrun.bold = True
            Usrun.font.name = "Microsoft YaHei Light"
            Usrun.font.size = Pt(10)
            now_paragraph.add_run("   ")
            Zhrun = now_paragraph.add_run(Zhword)
            Zhrun.font.name = "Microsoft YaHei Light"
            Zhrun.element.rPr.rFonts.set(qn("w:eastAsia"), "微软雅黑")
            Zhrun.font.size = Pt(10)
            now_paragraph.paragraph_format.space_after = Pt(2)

    doc.save("短语总结[第%s版 %s].docx" %(counter, func_word))
    Noticer("Phrases Exporter", "导出完毕，可进行下一步")
    return
def Ex_sorter(mode):
    if not Asker("Phrases Sorter", "是否确认排序？"): return
    #print("sorting")
    list = Files_finder()
    for files in list:
        if mode: Sort_phrases(files)
        else: Shuffle_phrases(files)
    #print("sorted")
    Noticer("Phrases Sorter", "排序完毕，可进行下一步")
    return

if __name__ == '__main__':
    Manager_Window = Tk(baseName = "0", className = "0")
    Manager_Window.title("Console")
    Manager_Window.resizable(False, False)
    Manager_Window.geometry("+570+390")

    Button(Manager_Window, text = "添加短语", width = 30, height = 5, command = lambda: {Manager_Window.destroy(), Phrases_Adder()}).grid(row = 0, column = 0, sticky = "w", padx = 10, pady = 5)
    Button(Manager_Window, text = "立即测试", width = 30, height = 5, command = lambda: {Manager_Window.destroy(), Phrases_Tester()}).grid(row = 0, column = 1, sticky = "e", padx = 10, pady = 5)
    Button(Manager_Window, text = "有序导出", width = 30, height = 5, command = lambda: {Manager_Window.destroy(), Exporter(True)}).grid(row = 1, column = 0, sticky = "e", padx = 10, pady = 5)
    Button(Manager_Window, text = "无序导出", width = 30, height = 5, command = lambda: {Manager_Window.destroy(), Exporter(False)}).grid(row = 1, column = 1, sticky = "e", padx = 10, pady = 5)
    Button(Manager_Window, text = "短语有序化", width = 30, height = 5, command = lambda: {Manager_Window.destroy(), Ex_sorter(True)}).grid(row = 2, column = 0, sticky = "w", padx = 10, pady = 5)
    Button(Manager_Window, text = "短语无序化", width = 30, height = 5, command = lambda: {Manager_Window.destroy(), Ex_sorter(False)}).grid(row = 2, column = 1, sticky = "e", padx = 10, pady = 5)

    Manager_Window.mainloop()